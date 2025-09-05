import os
import shutil
import re
from docx import Document
from openpyxl import load_workbook

TEMPLATE_FILE = "certificate.docx"
EXCEL_FILE = "internship_details.xlsx"
OUTPUT_DIR = "certificates"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Load workbook
wb = load_workbook(EXCEL_FILE, data_only=True)
sheet = wb.active

# Read headers from first row
headers = [str(cell.value).strip() if cell.value is not None else "" for cell in sheet[1]]

# Try common names for serial column; fallback to first column
try:
    slno_index = headers.index("Sl.No")
except ValueError:
    try:
        slno_index = headers.index("Sl No")
    except ValueError:
        slno_index = 0

# Placeholder regex: captures text inside { ... }, allowing spaces
PLACEHOLDER_RE = re.compile(r'\{\s*([^}]+?)\s*\}')

def copy_formatting(target_run, source_run):
    """Copy simple formatting from source run to target run (if available)."""
    try:
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
    except Exception:
        pass
    try:
        if source_run.font.name:
            target_run.font.name = source_run.font.name
    except Exception:
        pass
    try:
        if source_run.font.size:
            target_run.font.size = source_run.font.size
    except Exception:
        pass
    try:
        if source_run.font.color and source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb
    except Exception:
        pass

def replace_placeholders_in_paragraph(paragraph, mapping):
    """Replace {Key} placeholders in paragraph (even across runs), bold the replaced values,
       and preserve formatting by copying formatting from the run where the placeholder starts."""
    # Quick checks
    if not paragraph.runs:
        return
    full_text = ''.join(run.text for run in paragraph.runs)
    if "{" not in full_text:
        return
    matches = list(PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return

    # Build run_map: for each character position, which run index it came from
    run_map = []
    for idx, run in enumerate(paragraph.runs):
        text = run.text or ""
        run_map.extend([idx] * len(text))

    # Build sequence of segments: ('text'|'replace', text, char_position_for_formatting)
    segments = []
    last = 0
    for m in matches:
        s, e = m.start(), m.end()
        key = m.group(1).strip()
        if s > last:
            segments.append(('text', full_text[last:s], last))
        if key in mapping:
            segments.append(('replace', mapping[key], s))
        else:
            # keep placeholder as-is if no mapping found
            segments.append(('text', full_text[s:e], s))
        last = e
    if last < len(full_text):
        segments.append(('text', full_text[last:], last))

    # Clear existing runs' text (keep run objects so we can copy formatting from them)
    for run in paragraph.runs:
        run.text = ''

    # Append new runs for segments
    for seg_type, seg_text, seg_pos in segments:
        # determine source run for formatting (choose the run at seg_pos or fallback to last run)
        if run_map:
            if seg_pos < len(run_map):
                src_idx = run_map[seg_pos]
            else:
                src_idx = run_map[-1]
            src_run = paragraph.runs[src_idx]
        else:
            src_run = None

        new_run = paragraph.add_run(seg_text)
        if src_run is not None:
            copy_formatting(new_run, src_run)
        # If this is a replaced value, force bold (user-requested)
        if seg_type == 'replace':
            new_run.bold = True

def process_document(doc, mapping):
    # paragraphs
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, mapping)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, mapping)

    # headers and footers (all sections)
    for section in doc.sections:
        # header
        header = section.header
        if header is not None:
            for p in header.paragraphs:
                replace_placeholders_in_paragraph(p, mapping)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_placeholders_in_paragraph(p, mapping)
        # footer
        footer = section.footer
        if footer is not None:
            for p in footer.paragraphs:
                replace_placeholders_in_paragraph(p, mapping)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_placeholders_in_paragraph(p, mapping)

# MAIN loop: iterate rows, build mapping, duplicate template, replace, save
for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
    try:
        sl = row[slno_index]
        if sl is None or str(sl).strip() == "":
            print(f"Skipping row {row_idx} (empty Sl.No)")
            continue
        slno = str(sl).strip()
        output_file = os.path.join(OUTPUT_DIR, f"{slno}.docx")

        # make a fresh copy of the template
        shutil.copy2(TEMPLATE_FILE, output_file)
        doc = Document(output_file)

        # Build mapping {Header: value}
        mapping = {}
        for i, h in enumerate(headers):
            key = str(h).strip() if h else ""
            val = row[i] if i < len(row) else None
            mapping[key] = "" if val is None else str(val)

        # Print mapping for debug/visibility
        print(f"\nRow {row_idx} -> Sl.No {slno} mapping:")
        for k, v in mapping.items():
            print(f"  {k} -> {v}")

        # Process replacements (paragraphs, tables, headers, footers)
        process_document(doc, mapping)

        # Save
        doc.save(output_file)
        print(f"Generated: {output_file}")

    except Exception as exc:
        print(f"Error processing row {row_idx}: {exc}")

print("\nDone.")
