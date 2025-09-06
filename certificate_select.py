import os
import shutil
import re
from datetime import datetime
from docx import Document
from openpyxl import load_workbook
import tkinter as tk
from tkinter import ttk, messagebox

# ---------------- CONFIG ----------------
TEMPLATE_FILENAME = "certificate.docx"
EXCEL_FILENAME = "internship_details.xlsx"
OUTPUT_DIRNAME = "certificates"

current_directory = os.getcwd()
template_path = os.path.join(current_directory, TEMPLATE_FILENAME)
excel_path = os.path.join(current_directory, EXCEL_FILENAME)
output_dir = os.path.join(current_directory, OUTPUT_DIRNAME)
os.makedirs(output_dir, exist_ok=True)

# ---------------- HELPERS ----------------
PLACEHOLDER_RE = re.compile(r'\{\s*([^}]+?)\s*\}')

def format_value(val):
    if val is None:
        return ""
    if isinstance(val, datetime):
        return val.strftime("%d-%b-%Y")
    if isinstance(val, str):
        cleaned = val.strip()
        try:
            c = (cleaned.replace("st","").replace("nd","")
                        .replace("rd","").replace("th",""))
            parsed = datetime.strptime(c.strip(), "%d %B %Y")
            return parsed.strftime("%d-%b-%Y")
        except:
            try:
                c = (cleaned.replace("st","").replace("nd","")
                            .replace("rd","").replace("th",""))
                parsed = datetime.strptime(c.strip(), "%d %b %Y")
                return parsed.strftime("%d-%b-%Y")
            except:
                return cleaned
    return str(val)

def copy_formatting(target_run, source_run):
    try: target_run.bold = source_run.bold
    except: pass
    try: target_run.italic = source_run.italic
    except: pass
    try: target_run.underline = source_run.underline
    except: pass
    try:
        if source_run.font.name:
            target_run.font.name = source_run.font.name
    except: pass
    try:
        if source_run.font.size:
            target_run.font.size = source_run.font.size
    except: pass
    try:
        if source_run.font.color and source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb
    except: pass

def replace_placeholders_in_paragraph(paragraph, mapping):
    if not paragraph.runs:
        return
    full_text = ''.join(run.text for run in paragraph.runs)
    if "{" not in full_text:
        return
    matches = list(PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return
    run_map = []
    for idx, run in enumerate(paragraph.runs):
        text = run.text or ""
        run_map.extend([idx] * len(text))
    segments = []
    last = 0
    for m in matches:
        s, e = m.start(), m.end()
        key = m.group(1).strip()
        if s > last:
            segments.append(("text", full_text[last:s], last))
        if key in mapping:
            segments.append(("replace", mapping[key], s))
        else:
            segments.append(("text", full_text[s:e], s))
        last = e
    if last < len(full_text):
        segments.append(("text", full_text[last:], last))
    for run in paragraph.runs:
        run.text = ""
    for seg_type, seg_text, seg_pos in segments:
        if run_map:
            src_idx = run_map[min(seg_pos, len(run_map) - 1)]
            src_run = paragraph.runs[src_idx]
        else:
            src_run = None
        new_run = paragraph.add_run(seg_text)
        if src_run:
            copy_formatting(new_run, src_run)
        if seg_type == "replace":
            new_run.bold = True

def process_document(doc, mapping):
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, mapping)
    for section in doc.sections:
        for container in (section.header, section.footer):
            if container:
                for p in container.paragraphs:
                    replace_placeholders_in_paragraph(p, mapping)
                for table in container.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                replace_placeholders_in_paragraph(p, mapping)

# ---------------- DATA LOADING ----------------
wb = load_workbook(excel_path, data_only=True)
sheet = wb.active
headers = [str(cell.value).strip() if cell.value else "" for cell in sheet[1]]

if "Certificate No." not in headers:
    raise ValueError("Excel must contain a column named 'Certificate No.'")

certno_index = headers.index("Certificate No.")
name_index = headers.index("name") if "name" in headers else None

cert_data = []
for row in sheet.iter_rows(min_row=2, values_only=True):
    cert_val = row[certno_index]
    if not cert_val:
        continue
    cert_no = str(cert_val).strip()
    name = ""
    if name_index is not None and row[name_index]:
        name = str(row[name_index]).strip()
    cert_data.append((cert_no, name, row))

# ---------------- GENERATION ----------------
def generate_certificates(records):
    count = 0
    for cert_no, name, row in records:
        output_file = os.path.join(output_dir, f"{cert_no}.docx")
        try:
            shutil.copy2(template_path, output_file)
            doc = Document(output_file)
        except Exception as e:
            messagebox.showerror("File error", f"Error with template for {cert_no}: {e}")
            continue
        mapping = {}
        for i, h in enumerate(headers):
            key = str(h).strip() if h else ""
            val = row[i] if i < len(row) else None
            mapping[key] = format_value(val)
        process_document(doc, mapping)
        try:
            doc.save(output_file)
            count += 1
        except Exception as e:
            messagebox.showerror("Save error", f"Could not save {output_file}: {e}")
    messagebox.showinfo("Done", f"Generated {count} certificate(s).")

# ---------------- UI ----------------
def on_mode_change(event=None):
    choice = mode_var.get()
    frame_single.pack_forget()
    frame_selected.pack_forget()
    frame_all.pack_forget()

    if choice == "Single Certificate":
        frame_single.pack(pady=6)
    elif choice == "Selected Certificates":
        frame_selected.pack(pady=6)
    elif choice == "All Certificates":
        frame_all.pack(pady=6)

def generate():
    choice = mode_var.get()
    if choice == "Single Certificate":
        cert_no = entry_single.get().strip()
        if not cert_no:
            messagebox.showwarning("Input needed", "Please enter a Certificate No.")
            return
        match = None
        for c_no, name, row in cert_data:
            if c_no == cert_no:
                match = (c_no, name, row)
                break
        if not match:
            messagebox.showerror("Not found", f"Certificate No. {cert_no} not found in Excel.")
            return
        generate_certificates([match])

    elif choice == "Selected Certificates":
        selected_indices = listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("No selection", "Please select at least one certificate.")
            return
        selected = [cert_data[int(idx)] for idx in selected_indices]
        generate_certificates(selected)

    elif choice == "All Certificates":
        generate_certificates(cert_data)

# ---------------- UI ----------------
root = tk.Tk()
root.title("Certificate Generator")
root.geometry("250x350")

# Mode selection
mode_var = tk.StringVar(value="Single Certificate")
tk.Label(root, text="Choose generation mode:").pack(pady=4)
mode_combo = ttk.Combobox(root, textvariable=mode_var, state="readonly",
                          values=["Single Certificate", "Selected Certificates", "All Certificates"])
mode_combo.pack(pady=4)
mode_combo.bind("<<ComboboxSelected>>", on_mode_change)

# --- Single Certificate Frame ---
frame_single = tk.Frame(root)
tk.Label(frame_single, text="Enter Certificate No.:").pack(pady=4)
entry_single = tk.Entry(frame_single, width=15)
entry_single.pack(pady=4)
btn_single = ttk.Button(frame_single, text="Generate Certificate", command=generate)
btn_single.pack(pady=6)

# --- Selected Certificates Frame ---
frame_selected = tk.Frame(root)
list_frame = tk.Frame(frame_selected)
list_frame.pack(pady=6, fill=tk.BOTH, expand=True)

scrollbar = tk.Scrollbar(list_frame, orient=tk.VERTICAL)
listbox = tk.Listbox(list_frame, selectmode=tk.MULTIPLE, width=30, height=8, yscrollcommand=scrollbar.set)
scrollbar.config(command=listbox.yview)
listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

# Populate listbox
for cert_no, name, _ in cert_data:
    display = f"{cert_no} - {name}" if name else cert_no
    listbox.insert(tk.END, display)

btn_selected = ttk.Button(frame_selected, text="Generate Certificates", command=generate)
btn_selected.pack(pady=6)

# --- All Certificates Frame ---
frame_all = tk.Frame(root)
btn_all = ttk.Button(frame_all, text="Generate All Certificates", command=generate)
btn_all.pack(pady=10)

# Initialize view
on_mode_change()
root.mainloop()
